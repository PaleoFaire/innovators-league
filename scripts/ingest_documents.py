#!/usr/bin/env python3
"""
Opti Phase 2 — Document Ingestion Script
Processes .docx and .pdf files, sends chunks to the opti-embed Edge Function
which generates gte-small embeddings and inserts into Supabase pgvector.

Usage:
  .venv/bin/python3 scripts/ingest_documents.py --all                # Process all sources
  .venv/bin/python3 scripts/ingest_documents.py --source podcast     # Process only podcasts
  .venv/bin/python3 scripts/ingest_documents.py --dry-run --all      # Preview without uploading
  .venv/bin/python3 scripts/ingest_documents.py --all --clear        # Clear existing + re-upload

Prerequisites:
  # Already installed in .venv:
  #   python-docx, PyPDF2, supabase, tiktoken, tqdm, requests
  export SUPABASE_SERVICE_ROLE_KEY='your-service-role-key-here'

Architecture:
  Text extraction + chunking happens locally (Python).
  Embedding generation + database insertion happens server-side
  via the opti-embed Supabase Edge Function (gte-small built-in).
"""

import os
import sys
import json
import re
import argparse
import time
from pathlib import Path
from typing import List, Dict, Optional

import requests

# ══════════════════════════════════════════════════
# CONFIGURATION
# ══════════════════════════════════════════════════

SUPABASE_URL = 'https://imxrdesbozbxmlffewyr.supabase.co'
SUPABASE_SERVICE_KEY = os.environ.get('SUPABASE_SERVICE_ROLE_KEY', '')
SUPABASE_ANON_KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImlteHJkZXNib3pieG1sZmZld3lyIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzE0NTM2MDQsImV4cCI6MjA4NzAyOTYwNH0.ULUcvwaU2d-gyL_k6lcRziIEsndnDLGtMkV61Q9Knq0'

# Edge Function URLs
EMBED_FUNCTION_URL = f'{SUPABASE_URL}/functions/v1/opti-embed'

# Chunking parameters
CHUNK_SIZE = 400        # Target tokens per chunk
CHUNK_OVERLAP = 50      # Overlap tokens between consecutive chunks
MAX_TOKENS = 480        # Hard ceiling (gte-small limit is 512, leave margin)

# How many chunks to send per Edge Function call
UPLOAD_BATCH_SIZE = 10  # Small batches since each needs embedding generation

# Document source directories
SOURCES = {
    'podcast': {
        'path': '/Users/stephenmcbride/Desktop/Podcast',
        'extensions': ['.docx'],
        'description': 'Podcast interview transcripts',
        'recursive': True,
    },
    'research': {
        'path': '/Users/stephenmcbride/Desktop/1DI',
        'extensions': ['.docx', '.pdf'],
        'description': 'Monthly research reports (1DI)',
        'recursive': True,
    },
    'notes': {
        'path': '/Users/stephenmcbride/Desktop',
        'extensions': ['.docx'],
        'description': 'Top-level research notes',
        'recursive': False,
        'explicit_files': [
            'AI_Disruption_Risk_Ranking_Top100_Feb2026.docx',
            'CA notes.docx',
            'Interview Research - Pablos Holman.docx',
            'Monday notes.docx',
            'Rational_Optimist_Chartbook_200_Curated.docx',
            'Zurich presentation.docx',
            'nick.docx',
            'ROS - Sebastian A. Brunemeier - Research and Questions.docx',
        ],
    },
    'optimism': {
        'path': '/Users/stephenmcbride/Desktop/Optimism',
        'extensions': ['.docx', '.pdf'],
        'description': 'Optimism-themed research and notes',
        'recursive': True,
    },
}

SKIP_PATTERNS = ['~$', '.DS_Store', '__pycache__', '.tmp']


# ══════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════

def extract_text_docx(filepath: str) -> Optional[str]:
    """Extract text from a .docx file."""
    from docx import Document as DocxDocument
    try:
        doc = DocxDocument(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return '\n\n'.join(paragraphs)
    except Exception as e:
        print(f"  [WARN] Failed to extract {filepath}: {e}")
        return None


def extract_text_pdf(filepath: str) -> Optional[str]:
    """Extract text from a .pdf file."""
    from PyPDF2 import PdfReader
    try:
        reader = PdfReader(filepath)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return '\n\n'.join(pages)
    except Exception as e:
        print(f"  [WARN] Failed to extract {filepath}: {e}")
        return None


def extract_text(filepath: str) -> Optional[str]:
    """Extract text from a file based on its extension."""
    ext = Path(filepath).suffix.lower()
    if ext == '.docx':
        return extract_text_docx(filepath)
    elif ext == '.pdf':
        return extract_text_pdf(filepath)
    return None


# ══════════════════════════════════════════════════
# CHUNKING
# ══════════════════════════════════════════════════

_tokenizer = None

def get_tokenizer():
    """Lazy-load tiktoken tokenizer."""
    global _tokenizer
    if _tokenizer is None:
        import tiktoken
        _tokenizer = tiktoken.get_encoding('cl100k_base')
    return _tokenizer


def count_tokens(text: str) -> int:
    """Count tokens (approximation for gte-small)."""
    return len(get_tokenizer().encode(text))


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE,
               overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks by token count."""
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    if not paragraphs:
        return []

    chunks = []
    current_paras = []
    current_tokens = 0

    for para in paragraphs:
        para_tokens = count_tokens(para)

        if para_tokens > MAX_TOKENS:
            if current_paras:
                chunks.append('\n\n'.join(current_paras))
                current_paras = []
                current_tokens = 0
            sentences = re.split(r'(?<=[.!?])\s+', para)
            for sent in sentences:
                sent_tokens = count_tokens(sent)
                if current_tokens + sent_tokens > chunk_size and current_paras:
                    chunks.append('\n\n'.join(current_paras))
                    overlap_text = current_paras[-1] if current_paras else ''
                    current_paras = [overlap_text] if overlap_text else []
                    current_tokens = count_tokens(overlap_text) if overlap_text else 0
                current_paras.append(sent)
                current_tokens += sent_tokens
            continue

        if current_tokens + para_tokens > chunk_size and current_paras:
            chunks.append('\n\n'.join(current_paras))
            overlap_paras = []
            overlap_count = 0
            for p in reversed(current_paras):
                p_tokens = count_tokens(p)
                if overlap_count + p_tokens > overlap:
                    break
                overlap_paras.insert(0, p)
                overlap_count += p_tokens
            current_paras = overlap_paras
            current_tokens = overlap_count

        current_paras.append(para)
        current_tokens += para_tokens

    if current_paras:
        chunks.append('\n\n'.join(current_paras))

    return chunks


# ══════════════════════════════════════════════════
# FILE DISCOVERY
# ══════════════════════════════════════════════════

def discover_files(source_key: str) -> List[str]:
    """Discover all processable files for a given source."""
    config = SOURCES[source_key]
    base_path = config['path']
    extensions = config['extensions']
    recursive = config.get('recursive', True)
    explicit_files = config.get('explicit_files', None)

    files = []

    if explicit_files:
        for fname in explicit_files:
            fpath = os.path.join(base_path, fname)
            if os.path.isfile(fpath):
                files.append(fpath)
            else:
                print(f"  [WARN] Explicit file not found: {fpath}")
    elif recursive:
        for root, dirs, fnames in os.walk(base_path):
            dirs[:] = [d for d in dirs if not d.startswith('.') and d != '__pycache__']
            for fname in fnames:
                if any(skip in fname for skip in SKIP_PATTERNS):
                    continue
                if Path(fname).suffix.lower() in extensions:
                    files.append(os.path.join(root, fname))
    else:
        if os.path.isdir(base_path):
            for fname in os.listdir(base_path):
                if any(skip in fname for skip in SKIP_PATTERNS):
                    continue
                if Path(fname).suffix.lower() in extensions:
                    fpath = os.path.join(base_path, fname)
                    if os.path.isfile(fpath):
                        files.append(fpath)

    return sorted(files)


# ══════════════════════════════════════════════════
# UPLOAD VIA EDGE FUNCTION (embed + insert)
# ══════════════════════════════════════════════════

def upload_via_edge_function(chunks_data: List[Dict], batch_size: int = UPLOAD_BATCH_SIZE) -> int:
    """Send chunks to the opti-embed Edge Function which generates embeddings and inserts."""
    from tqdm import tqdm

    uploaded = 0
    total_batches = (len(chunks_data) + batch_size - 1) // batch_size

    for i in tqdm(range(0, len(chunks_data), batch_size), desc="Uploading", total=total_batches):
        batch = chunks_data[i:i + batch_size]

        # Prepare payload for the Edge Function
        payload = {
            'chunks': [
                {
                    'content': c['content'],
                    'source_file': c['source_file'],
                    'source_type': c['source_type'],
                    'chunk_index': c['chunk_index'],
                    'token_count': c['token_count'],
                    'metadata': c['metadata'],
                }
                for c in batch
            ]
        }

        try:
            resp = requests.post(
                EMBED_FUNCTION_URL,
                headers={
                    'Content-Type': 'application/json',
                    'apikey': SUPABASE_ANON_KEY,
                    'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
                },
                json=payload,
                timeout=120,  # Embedding generation can take time
            )

            if resp.status_code == 200:
                result = resp.json()
                uploaded += result.get('inserted', 0)
                errors = result.get('errors', [])
                if errors:
                    for err in errors[:3]:
                        print(f"  [WARN] {err}")
            else:
                print(f"  [ERROR] Batch failed (HTTP {resp.status_code}): {resp.text[:200]}")
                # Retry individual chunks
                for c in batch:
                    try:
                        single_resp = requests.post(
                            EMBED_FUNCTION_URL,
                            headers={
                                'Content-Type': 'application/json',
                                'apikey': SUPABASE_ANON_KEY,
                                'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
                            },
                            json={'chunks': [c]},
                            timeout=60,
                        )
                        if single_resp.status_code == 200:
                            uploaded += single_resp.json().get('inserted', 0)
                    except Exception:
                        pass

        except requests.exceptions.Timeout:
            print(f"  [WARN] Batch timed out, retrying individually...")
            for c in batch:
                try:
                    single_resp = requests.post(
                        EMBED_FUNCTION_URL,
                        headers={
                            'Content-Type': 'application/json',
                            'apikey': SUPABASE_ANON_KEY,
                            'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
                        },
                        json={'chunks': [c]},
                        timeout=60,
                    )
                    if single_resp.status_code == 200:
                        uploaded += single_resp.json().get('inserted', 0)
                except Exception:
                    pass
        except Exception as e:
            print(f"  [ERROR] Batch exception: {e}")

        # Small delay between batches to avoid rate limits
        time.sleep(0.5)

    return uploaded


def clear_source_via_supabase(source_type: str):
    """Delete existing documents for a source type via Supabase client."""
    from supabase import create_client
    client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    try:
        client.table('opti_documents').delete().eq('source_type', source_type).execute()
        print(f"  Cleared existing '{source_type}' documents.")
    except Exception as e:
        print(f"  [WARN] Failed to clear source '{source_type}': {e}")


# ══════════════════════════════════════════════════
# MAIN PIPELINE
# ══════════════════════════════════════════════════

def process_source(source_key: str, dry_run: bool = False,
                   clear_existing: bool = False) -> Dict:
    """Process all files for a source and upload via Edge Function."""
    from tqdm import tqdm

    config = SOURCES[source_key]
    print(f"\n{'=' * 60}")
    print(f"Processing: {config['description']} ({source_key})")
    print(f"Path: {config['path']}")
    print(f"{'=' * 60}")

    files = discover_files(source_key)
    print(f"Found {len(files)} files to process.")

    if not files:
        return {'source': source_key, 'files': 0, 'chunks': 0, 'uploaded': 0}

    if clear_existing and not dry_run:
        clear_source_via_supabase(source_key)

    all_chunks_data = []
    skipped = 0

    for filepath in tqdm(files, desc=f"Extracting {source_key}"):
        filename = os.path.basename(filepath)

        file_size = os.path.getsize(filepath)
        if file_size > 50 * 1024 * 1024:
            print(f"  [SKIP] {filename} — too large ({file_size // (1024*1024)}MB)")
            skipped += 1
            continue

        text = extract_text(filepath)
        if not text or len(text.strip()) < 100:
            skipped += 1
            continue

        chunks = chunk_text(text)

        metadata = {
            'title': Path(filename).stem,
            'file_size': file_size,
        }

        for idx, chunk in enumerate(chunks):
            all_chunks_data.append({
                'content': chunk,
                'source_file': filename,
                'source_type': source_key,
                'chunk_index': idx,
                'token_count': count_tokens(chunk),
                'metadata': metadata,
            })

    print(f"Generated {len(all_chunks_data)} chunks from {len(files) - skipped} files ({skipped} skipped).")

    if not all_chunks_data:
        return {'source': source_key, 'files': len(files), 'chunks': 0, 'uploaded': 0}

    if dry_run:
        print(f"[DRY RUN] Would upload {len(all_chunks_data)} chunks via Edge Function.")
        for chunk in all_chunks_data[:3]:
            print(f"  File: {chunk['source_file']}, Chunk {chunk['chunk_index']}, "
                  f"Tokens: {chunk['token_count']}")
            print(f"  Preview: {chunk['content'][:120]}...")
        return {'source': source_key, 'files': len(files),
                'chunks': len(all_chunks_data), 'uploaded': 0}

    # Upload via Edge Function (which generates embeddings server-side)
    print("Uploading to Supabase via Edge Function (embedding + insert)...")
    uploaded = upload_via_edge_function(all_chunks_data)
    print(f"Uploaded {uploaded}/{len(all_chunks_data)} chunks.")

    return {'source': source_key, 'files': len(files),
            'chunks': len(all_chunks_data), 'uploaded': uploaded}


def main():
    parser = argparse.ArgumentParser(description='Opti Phase 2 — Document Ingestion')
    parser.add_argument('--all', action='store_true', help='Process all sources')
    parser.add_argument('--source', type=str, choices=list(SOURCES.keys()),
                        help='Process a specific source')
    parser.add_argument('--dry-run', action='store_true',
                        help='Extract and chunk without uploading')
    parser.add_argument('--clear', action='store_true',
                        help='Clear existing documents before uploading')
    args = parser.parse_args()

    if not args.all and not args.source:
        parser.print_help()
        sys.exit(1)

    if not args.dry_run and not SUPABASE_SERVICE_KEY:
        print("\n[ERROR] Set SUPABASE_SERVICE_ROLE_KEY environment variable.")
        print("  Find it at: Supabase Dashboard > Project Settings > API > service_role secret")
        print("  Then run:")
        print("    export SUPABASE_SERVICE_ROLE_KEY='your-key-here'")
        sys.exit(1)

    results = []

    if args.all:
        for source_key in SOURCES:
            result = process_source(source_key, args.dry_run, args.clear)
            results.append(result)
    elif args.source:
        result = process_source(args.source, args.dry_run, args.clear)
        results.append(result)

    # Summary
    print(f"\n{'=' * 60}")
    print("INGESTION SUMMARY")
    print(f"{'=' * 60}")
    total_files = 0
    total_chunks = 0
    total_uploaded = 0
    for r in results:
        print(f"  {r['source']:12s} — {r['files']:3d} files, "
              f"{r['chunks']:5d} chunks, {r['uploaded']:5d} uploaded")
        total_files += r['files']
        total_chunks += r['chunks']
        total_uploaded += r['uploaded']
    print(f"  {'TOTAL':12s} — {total_files:3d} files, "
          f"{total_chunks:5d} chunks, {total_uploaded:5d} uploaded")


if __name__ == '__main__':
    main()
